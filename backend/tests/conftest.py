import io

import fitz
import pytest
from docx import Document


@pytest.fixture
def text_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text(
        (72, 72),
        "Jane Doe\njane@example.com\nSoftware Engineer at Acme Corp",
    )
    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


@pytest.fixture
def scanned_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page()
    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


@pytest.fixture
def password_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Protected resume content")
    buffer = io.BytesIO()
    document.save(
        buffer,
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw="secret",
    )
    document.close()
    return buffer.getvalue()


@pytest.fixture
def docx_with_table_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Header paragraph")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Table cell A"
    table.rows[0].cells[1].text = "Table cell B"
    document.add_paragraph("Footer paragraph")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def utf8_txt_bytes() -> bytes:
    return "Name: Alex\nRole: Engineer\n\nSkills: Python, FastAPI".encode("utf-8")


@pytest.fixture
def sample_resume() -> "ResumeSchema":
    from app.models.domain.resume import ContactSchema, ResumeSchema, SkillsSchema, WorkExperienceEntry

    return ResumeSchema(
        contact=ContactSchema(name="Jane Doe", email="jane@example.com"),
        summary="Backend engineer with API experience.",
        work_experience=[
            WorkExperienceEntry(
                company="Acme Corp",
                title="Software Engineer",
                duration="2021 - Present",
                bullets=["Built internal APIs", "Improved deploy speed"],
            )
        ],
        skills=SkillsSchema(languages=["Python"], frameworks=["FastAPI"]),
        projects=[],
    )


@pytest.fixture
def sample_jd() -> "JDAnalysisSchema":
    from app.models.domain.jd import JDAnalysisSchema

    return JDAnalysisSchema(
        role_title="Senior Backend Engineer",
        must_have_skills=["Python", "FastAPI"],
        nice_to_have_skills=["PostgreSQL"],
        keywords=["API", "microservices"],
        seniority="senior",
        tone="technical",
        responsibilities=["Design scalable APIs"],
    )
